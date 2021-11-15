from docx import Document
import json 


#Steps to run this code: 
    # 1. import docx and json using (pip)or pip3 install python-docx
    # 2. run python script.py in the same directory or update the path and it will output a data.json file 
    # 3. right now the configuration is that the last 4 tables are not being used and if they ever want to be used, just add the header to the key and get rid of a pop() statement corresponding to that key
    # 4. **IMPORTANT** make sure that the language and en variables in line 40-41 match up to table names


#If you want to iterate through directory with multiple docx files, edit this piece of code

# import os
# for root, dirs, files in os.walk('/path/to/your-directory'):
#     for file in files:
#         if file.endswith('.docx'):
#             print(os.path.join(root, file))

document = Document('russian.docx')
# document = Document('farsi.docx')

initial_dict = {}
#if ever you need to add mroe information to the headers and tables just update this
headers_key = ["header", "vaccinepage", "dashboardpage", "vaccineform", "qrpage", "receivedpage", "footer", "remove1", "remove2", "remove3", "remove4"] 

for i in range(len(document.tables)): 
    table = document.tables[i]
    keys = None
    data = []
    row_data = None
    for j, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if j == 0: 
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
        initial_dict[i] = data


initial_dict.pop(0) #unnecessary info
initial_dict.pop(1) #unnecessary info

overall_dict = {}


#EDIT THIS TO MATCH TABLE HEADERS
language = 'Language'
en = 'en'


# print(output_dict.get(2)) # <- run this to make sure that the output is correct in the hashmap - error handling


for ind, val in initial_dict.items():
    temp_dict = {}
    for each_entry in val: 
        temp_dict[each_entry.get(language)] = each_entry.get(en) #make sure these line up with the EDIT section
    overall_dict[headers_key[ind-2]] = temp_dict

#removes the unncessary tables (but has the information if ever needed to use them)
overall_dict.pop('remove4')
overall_dict.pop('remove3')
overall_dict.pop('remove2')
overall_dict.pop('remove1')
# print (overall_dict)

jsonData = json.dumps(overall_dict, indent=4, separators=(", ", ": "), ensure_ascii=False).replace('null', '""')

# print (jsonData)

jsonFile = open('data.json', 'w')
jsonFile.write(jsonData)