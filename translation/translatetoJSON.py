from docx import Document
import json 
import os
from faq import parseFaq


#Steps to run this code: 
    # 1. import docx and json using (pip)or pip3 install python-docx
    # 2. run python script.py and it will output a data.json file 
    # 3. right now the configuration is that the last 4 tables are not being used and if they ever want to be used, just add the header to the key and get rid of a pop() statement corresponding to that key
    # 4. **IMPORTANT** make sure that the language header matches up to table name


# document = Document('russian.docx')
# document = Document('farsi.docx')
# document = Document('Spanish.docx')

path = './translate'
folder = os.fsencode(path)

documents = []

for file in os.listdir(folder):
    document = os.fsdecode(file)
    if document.endswith( ('.docx') ):
        documents.append(document)

print(documents)


for z in documents: 

    document = Document('./translate/' + str(z))


    initial_dict = {}
    #if ever you need to add mroe information to the headers and tables just update this
    headers_key = ["header", "vaccinepage", "dashboardpage", "vaccineform", "qrpage", "receivedpage", "footer", "remove1", "remove2", "remove3", "remove4"] 

    for i in range(len(document.tables)): 
        table = document.tables[i]
        keys = None
        data = []
        row_data = None
        for j, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            if j == 0: 
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            data.append(row_data)
            initial_dict[i] = data


    initial_dict.pop(0) #unnecessary info
    initial_dict.pop(1) #unnecessary info

    overall_dict = {}


    #EDIT THIS TO MATCH TABLE HEADERS
    language = 'Language'


    # print(output_dict.get(2)) # <- run this to make sure that the output is correct in the hashmap - error handling


    for ind, val in initial_dict.items():
        temp_dict = {}
        for each_entry in val: 
            temp_dict[each_entry.get(language)] = each_entry.get('En') or each_entry.get('en') #make sure these line up with the EDIT section
        overall_dict[headers_key[ind-2]] = temp_dict

    #removes the unncessary tables (but has the information if ever needed to use them)
    overall_dict.pop('remove4')
    overall_dict.pop('remove3')
    overall_dict.pop('remove2')
    overall_dict.pop('remove1')
    overall_dict.pop('qrpage') #removed cause it's empty
    # print (overall_dict)

    # Add faqpage
    parseFaq(document, overall_dict)

    jsonData = json.dumps(overall_dict, indent=4, separators=(", ", ": "), ensure_ascii=False).replace('null', '""')

    filename = "output" + str(z) + ".json"
    jsonFile = open(filename, 'w')
    jsonFile.write(jsonData)