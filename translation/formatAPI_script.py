from docx import Document
import json
import os
import re

# Steps to run this code: 
    # 1. import docx and json using (pip)or pip3 install python-docx
    # 2. Create a directory called translate and put all the word docs that you want formatted
    # 3. run python3 formatAPI_script.py
    # 4. files will be output for each doc

path = './translate'
folder = os.fsencode(path)

documents = []

for file in os.listdir(folder):
    document = os.fsdecode(file)
    if document.endswith( ('.docx') ):
        documents.append(document)

for z in documents:
    document = Document('./translate/' + str(z))

    initial_dict = {}
    #if ever you need to add mroe information to the headers and tables just update this
    headers_key = ["header", "vaccinepage", "dashboardpage", "vaccineform", "qrpage", "receivedpage", "footer", "FormatSms", "FormatNotFoundSms", "FormatHTML", "FormatNotFoundHTML"]

    for i in range(len(document.tables)):
        table = document.tables[i]
        keys = None
        data = []
        row_data = None
        for j, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            if j == 0: 
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            data.append(row_data)
            initial_dict[i] = data


    initial_dict.pop(0) #unnecessary info
    initial_dict.pop(1) #unnecessary info

    overall_dict = {}


    #EDIT THIS TO MATCH TABLE HEADERS
    language = 'Language'

    # print(output_dict.get(2)) # <- run this to make sure that the output is correct in the hashmap - error handling


    for ind, val in initial_dict.items():
        temp_dict = {}
        for each_entry in val: 
            temp_dict[each_entry.get(language)] = each_entry.get('en') or each_entry.get('En') or each_entry.get('')#make sure these line up with the EDIT section
        overall_dict[headers_key[ind-2]] = temp_dict

    #removes the unncessary tables (but has the information if ever needed to use them)
    #formatHTML
    headinginfo_format = overall_dict.get('FormatHTML').get('heading')
    infoText_format = overall_dict.get('FormatHTML').get('infoText').replace('24', '{linkExpireHours}')

    viewLink_format = overall_dict.get('FormatHTML').get('viewlink') or overall_dict.get('FormatHTML').get('mmatlink')
    viewLink_format = viewLink_format.replace('<a>', '<a href=\'{url}\'>')

    learnMore_format = overall_dict.get('FormatHTML').get('learnMore')
    learnMore_format = learnMore_format.replace('<a>', '<a href=\'{cdcUrl}\'>')

    questions_format = overall_dict.get('FormatHTML').get('questions')

    visitFAQ_format = overall_dict.get('FormatHTML').get('visitFAQ')
    visitFAQ_format = visitFAQ_format.replace('<a>', '<a href=\'{vaccineFAQUrl}\'>')

    stayInformed_format = overall_dict.get('FormatHTML').get('stayInformed')

    viewInfo_format = overall_dict.get('FormatHTML').get('viewInfo') or overall_dict.get('FormatHTML').get('viewnfo')
    viewInfo_format = viewInfo_format.replace('<a>', '<a href=\'{covidWebUrl}\'>')

    emailLabel_format = overall_dict.get('FormatHTML').get('emailLabel')

    output_str = '----------------------FORMATHTML----------------------------- \n\n\n'

    output_str += '$"<img src=\'{webUrl}/imgs/MyTurn-logo.png\'><br/>" +\n'
    output_str += '$"<h3 style=\'color: #f06724\'>' + headinginfo_format + '</h3>" +\n'
    output_str +=  '$"<p>' + infoText_format + '</p>" +\n'
    output_str += '$"<p><a href=\'{url}\'>' + viewLink_format + '</a></p>" +\n' if not '</a>' in viewLink_format else '$"<p>' + viewLink_format + '</p>" +\n'
    output_str += '$"<p>' + learnMore_format + '</p>" +\n'
    output_str += '$"<p><b>' + questions_format + '</b></p>" +\n'
    output_str += '$"<p>' + visitFAQ_format + '</p>" +\n'
    output_str += '$"<p><b>' + stayInformed_format + '</b></p>" +\n'
    output_str += '$"<p>' + viewInfo_format + '</p><br/>" +\n'
    output_str += '$"<hr>" +\n'
    output_str += '$"<footer><p style=\'text-align:center\'>' + emailLabel_format + '</p>" +\n'
    output_str += '$"<p style=\'text-align:center\'><img src=\'{emailLogoUrl}\'></p></footer>",'


    #htmlnotfoundformat

    output_str += '\n ----------------------------------FormatNotFoundHTML-------------------\n\n\n'

    headinginfo_noformat = overall_dict.get('FormatNotFoundHTML').get('heading')

    infoText_noformat = overall_dict.get('FormatNotFoundHTML').get('infoText')
    infoText_noformat = infoText_noformat.replace('<a>', '<a href=\'{webUrl}\'>')

    nextSteps_noformat = overall_dict.get('FormatNotFoundHTML').get('nextSteps')
    nextSteps_noformat = nextSteps_noformat.replace('<a>', '<a href=\'{webUrl}\'>', 1)
    nextSteps_noformat = nextSteps_noformat.replace('<a>', '<a href=\'{contactUsUrl}\'>')

    questions_noformat = overall_dict.get('FormatNotFoundHTML').get('questions')

    visitFAQ_noformat = overall_dict.get('FormatNotFoundHTML').get('visitFAQ')
    visitFAQ_noformat = visitFAQ_noformat.replace('<a>', '<a href=\'{vaccineFAQUrl}\'>').replace('\n', '')

    stayInformed_noformat = overall_dict.get('FormatNotFoundHTML').get('stayInformed')

    viewInfo_noformat = overall_dict.get('FormatNotFoundHTML').get('viewInfo') or overall_dict.get('FormatNotFoundHTML').get('mmatInfo')
    viewInfo_noformat = viewInfo_noformat.replace('<a>', '<a href=\'{covidWebUrl}\'>')

    emailLabel_noformat = overall_dict.get('FormatNotFoundHTML').get('emailLabel')

    HTMLnotfound_str = '$"<img src=\'{webUrl}/imgs/MyTurn-logo.png\'><br/>" +\n'
    HTMLnotfound_str += '$"<h3 style=\'color: #f06724\'>' + headinginfo_noformat + '</h3>" +\n'
    HTMLnotfound_str +=  '$"<p>' + infoText_noformat + '</p><br/>" +\n'
    HTMLnotfound_str += '$"<p>' + nextSteps_noformat + '</p>" +\n'
    HTMLnotfound_str += '$"<p><b>' + questions_noformat + '</b></p>" +\n'
    HTMLnotfound_str += '$"<p>' + visitFAQ_noformat + '</p>" +\n'
    HTMLnotfound_str += '$"<p><b>' + stayInformed_noformat + '</b></p>" +\n'
    HTMLnotfound_str += '$"<p>' + viewInfo_noformat + '</p><br/>" +\n'
    HTMLnotfound_str += '$"<hr>" +\n'
    HTMLnotfound_str += '$"<footer><p style=\'text-align:center\'>' + emailLabel_noformat + '</p>" +\n'
    HTMLnotfound_str += '$"<p style=\'text-align:center\'><img src=\'{emailLogoUrl}\'></p></footer>",'

    output_str += '\n ----------------------------------tagging output for HTMLNOTFOUND-------------------\n\n\n'

    output_str += HTMLnotfound_str

    # FormatSms string transform

    output_str += '\n ----------------------------------FormatSMS-------------------\n\n\n'

    formatsms_text = overall_dict.get('FormatSms').get('Text')
    formatsms_string = '$"{}"'.format(formatsms_text.replace("24", "{linkExpireHours}"))

    output_str += '\n ----------------------------------tagging output for formatSMS-------------------\n\n\n'

    output_str += formatsms_string

    # FormatNotFoundSms string transform

    output_str += '\n ----------------------------------FormatNotFoundSMS-------------------\n\n\n'


    formatnotfoundsms_text = overall_dict.get('FormatNotFoundSms').get('Text')
    formatnotfoundsms_string = '$"{}"'.format(re.sub(r'((1-\d{3}-\d{3}-\d{4})|(\(\d{3}\) \d{3}-\d{4})|(\d{3}-\d{3}-\d{4}))', "{phoneNumber}", formatnotfoundsms_text))
    
    output_str += '\n ----------------------------------tagging output for formatnotfoundSMS-------------------\n\n\n'

    output_str += formatnotfoundsms_string
    

    filename = "output" + str(z) + ".txt"
    strFile = open(filename, 'w')
    strFile.write(output_str)
